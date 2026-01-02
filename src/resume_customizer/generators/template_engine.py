"""Template rendering engine for resume document generation."""

import logging
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any

import jinja2
from weasyprint import HTML

from ..core.models import CustomizedResume, Skill, UserProfile

logger = logging.getLogger(__name__)


class TemplateNotFoundError(Exception):
    """Raised when a requested template cannot be found."""

    pass


class TemplateRenderError(Exception):
    """Raised when template rendering fails."""

    pass


class PDFGenerationError(Exception):
    """Raised when PDF generation fails."""

    pass


def _format_date_range(start_date: str, end_date: str) -> str:
    """
    Format date range for display.

    Args:
        start_date: Start date (e.g., "2020-01")
        end_date: End date (e.g., "2023-06" or "Present")

    Returns:
        Formatted date range (e.g., "Jan 2020 - Jun 2023" or "Jan 2020 - Present")
    """
    def format_date(date_str: str) -> str:
        """Format a single date."""
        if not date_str or date_str.lower() == "present":
            return "Present"

        try:
            # Try parsing YYYY-MM format
            if len(date_str) == 7 and date_str[4] == "-":
                dt = datetime.strptime(date_str, "%Y-%m")
                return dt.strftime("%b %Y")
            # Try parsing YYYY format
            elif len(date_str) == 4:
                return date_str
            # Return as-is if format is unknown
            else:
                return date_str
        except ValueError:
            return date_str

    start_formatted = format_date(start_date)
    end_formatted = format_date(end_date)

    return f"{start_formatted} - {end_formatted}"


def _group_skills_by_category(skills: list[Skill]) -> dict[str, list[dict[str, Any]]]:
    """
    Group skills by category.

    Args:
        skills: List of Skill objects

    Returns:
        Dictionary mapping category names to lists of skill dictionaries
    """
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)

    for skill in skills:
        skill_dict = {
            "name": skill.name,
            "proficiency": skill.proficiency,
            "years": skill.years,
            "description": skill.description,
        }
        grouped[skill.category].append(skill_dict)

    return dict(grouped)


def _prepare_template_context(
    customized_resume: CustomizedResume,
    user_profile: UserProfile,
) -> dict[str, Any]:
    """
    Prepare template context by merging CustomizedResume and UserProfile data.

    Args:
        customized_resume: Customized resume with selected/reordered data
        user_profile: Original user profile with complete information

    Returns:
        Dictionary with all template variables
    """
    # Prepare experiences with formatted dates and achievements
    experiences = []
    for exp in customized_resume.selected_experiences:
        exp_dict = {
            "company": exp.company,
            "title": exp.title,
            "location": exp.location,
            "work_mode": exp.work_mode,
            "dates": _format_date_range(exp.start_date, exp.end_date),
            "achievements": [
                {
                    "text": ach.text,
                    "technologies": ach.technologies,
                    "metrics": ach.metrics,
                }
                for ach in exp.achievements
            ],
            "technologies": exp.technologies,
        }
        experiences.append(exp_dict)

    # Prepare skills
    skills = [
        {
            "name": skill.name,
            "category": skill.category,
            "proficiency": skill.proficiency,
            "years": skill.years,
        }
        for skill in customized_resume.reordered_skills
    ]

    # Group skills by category
    skills_by_category = _group_skills_by_category(customized_resume.reordered_skills)

    # Prepare education
    education = [
        {
            "degree": edu.degree,
            "institution": edu.institution,
            "graduation_year": edu.graduation_year,
            "gpa": edu.gpa,
            "location": edu.location,
            "details": edu.details,
        }
        for edu in user_profile.education
    ]

    # Prepare certifications (optional)
    certifications = (
        [
            {
                "name": cert.name,
                "issuer": cert.issuer,
                "date": cert.date,
                "credential_id": cert.credential_id,
                "url": cert.url,
            }
            for cert in user_profile.certifications
        ]
        if user_profile.certifications
        else []
    )

    # Prepare projects (optional)
    projects = (
        [
            {
                "name": proj.name,
                "description": proj.description,
                "technologies": proj.technologies,
                "url": proj.url,
                "github": proj.github,
                "highlights": proj.highlights,
            }
            for proj in user_profile.projects
        ]
        if user_profile.projects
        else []
    )

    # Build complete context
    context = {
        # Personal information (from UserProfile)
        "name": user_profile.name,
        "contact": user_profile.contact.to_dict(),
        # Summary (customized if available)
        "summary": customized_resume.customized_summary or user_profile.summary,
        # Customized sections (from CustomizedResume)
        "experiences": experiences,
        "skills": skills,
        "skills_by_category": skills_by_category,
        # Unchanged sections (from UserProfile)
        "education": education,
        "certifications": certifications if certifications else None,
        "projects": projects if projects else None,
        # Metadata
        "metadata": {
            "match_score": customized_resume.match_result.overall_score,
            "customization_id": customized_resume.customization_id,
            "created_at": customized_resume.created_at,
            "template": customized_resume.template,
        },
    }

    return context


class TemplateEngine:
    """Jinja2-based template rendering engine for resume generation."""

    def __init__(self, templates_dir: Path | None = None):
        """
        Initialize template engine.

        Args:
            templates_dir: Custom templates directory (defaults to config.templates_dir)

        Raises:
            FileNotFoundError: If templates directory doesn't exist
        """
        if templates_dir is None:
            # Try to load from config
            try:
                from ..config import get_config

                config = get_config()
                templates_dir = config.templates_dir
            except Exception:
                # Fallback: use package-relative path
                package_root = Path(__file__).parent.parent.parent.parent
                templates_dir = package_root / "templates"

        if not templates_dir.exists():
            raise FileNotFoundError(
                f"Templates directory not found: {templates_dir}\n"
                f"Please ensure the templates/ directory exists and contains template files."
            )

        self.templates_dir = templates_dir
        logger.info(f"Template engine initialized with templates_dir: {templates_dir}")

        # Configure Jinja2 environment
        self.env = jinja2.Environment(
            loader=jinja2.FileSystemLoader(str(templates_dir)),
            autoescape=jinja2.select_autoescape(["html", "xml"]),
            trim_blocks=True,
            lstrip_blocks=True,
        )

        # Add custom filters
        self.env.filters["format_date_range"] = _format_date_range

    def list_templates(self) -> list[str]:
        """
        List available template names.

        Returns:
            List of template names (without .html extension)
        """
        template_files = list(self.templates_dir.glob("*.html"))
        return [t.stem for t in template_files]

    def load_template(self, template_name: str) -> jinja2.Template:
        """
        Load template by name.

        Args:
            template_name: Template name (e.g., "modern", "classic", "ats_optimized")

        Returns:
            Jinja2 Template object

        Raises:
            TemplateNotFoundError: If template doesn't exist
        """
        template_file = f"{template_name}.html"

        try:
            template = self.env.get_template(template_file)
            logger.debug(f"Loaded template: {template_file}")
            return template
        except jinja2.TemplateNotFound as e:
            available = self.list_templates()
            raise TemplateNotFoundError(
                f"Template '{template_name}' not found.\n"
                f"Available templates: {', '.join(available)}"
            ) from e

    def render_html(
        self,
        customized_resume: CustomizedResume,
        user_profile: UserProfile,
        template_name: str = "modern",
    ) -> str:
        """
        Render HTML from CustomizedResume and UserProfile.

        Args:
            customized_resume: The customized resume data
            user_profile: Original user profile (for contact, education, etc.)
            template_name: Template to use (modern/classic/ats_optimized)

        Returns:
            Rendered HTML string

        Raises:
            TemplateNotFoundError: If template doesn't exist
            TemplateRenderError: If rendering fails
        """
        try:
            # Load template
            template = self.load_template(template_name)

            # Prepare context
            context = _prepare_template_context(customized_resume, user_profile)

            # Render template
            html = template.render(**context)

            logger.info(
                f"Rendered HTML using template '{template_name}' "
                f"({len(html)} characters)"
            )

            return html

        except TemplateNotFoundError:
            raise
        except Exception as e:
            raise TemplateRenderError(
                f"Failed to render template '{template_name}': {e}"
            ) from e

    def generate_pdf(
        self,
        customized_resume: CustomizedResume,
        user_profile: UserProfile,
        output_path: Path,
        template_name: str = "modern",
    ) -> Path:
        """
        Generate PDF from CustomizedResume.

        Args:
            customized_resume: The customized resume data
            user_profile: Original user profile
            output_path: Path to save PDF
            template_name: Template to use

        Returns:
            Path to generated PDF file

        Raises:
            TemplateNotFoundError: If template doesn't exist
            PDFGenerationError: If PDF generation fails
        """
        try:
            # Render HTML
            html_content = self.render_html(
                customized_resume, user_profile, template_name
            )

            # Create output directory if needed
            output_path.parent.mkdir(parents=True, exist_ok=True)

            # Configure WeasyPrint
            base_url = str(self.templates_dir)
            html_obj = HTML(string=html_content, base_url=base_url)

            # Generate PDF
            pdf_bytes = html_obj.write_pdf()

            # Write to file
            output_path.write_bytes(pdf_bytes)

            logger.info(
                f"Generated PDF: {output_path} ({len(pdf_bytes)} bytes) "
                f"using template '{template_name}'"
            )

            return output_path

        except TemplateNotFoundError:
            raise
        except TemplateRenderError:
            raise
        except Exception as e:
            raise PDFGenerationError(
                f"Failed to generate PDF: {e}"
            ) from e

    def generate_docx(
        self,
        customized_resume: CustomizedResume,
        user_profile: UserProfile,
        output_path: Path,
        template_name: str = "modern",
    ) -> Path:
        """
        Generate DOCX from CustomizedResume.

        Args:
            customized_resume: The customized resume data
            user_profile: Original user profile
            output_path: Path to save DOCX
            template_name: Template style to use

        Returns:
            Path to generated DOCX file

        Raises:
            Exception: If DOCX generation fails
        """
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor

        try:
            # Prepare context
            context = _prepare_template_context(customized_resume, user_profile)

            # Create document
            doc = Document()

            # Configure margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)

            # Add name
            name_para = doc.add_paragraph(context["name"])
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_run = name_para.runs[0]
            name_run.font.size = Pt(20)
            name_run.font.bold = True
            if template_name == "modern":
                name_run.font.color.rgb = RGBColor(30, 64, 175)

            # Add contact info
            contact = context["contact"]
            contact_parts = []
            if contact.get("email"):
                contact_parts.append(contact["email"])
            if contact.get("phone"):
                contact_parts.append(contact["phone"])
            if contact.get("location"):
                contact_parts.append(contact["location"])

            contact_para = doc.add_paragraph(" | ".join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.runs[0]
            contact_run.font.size = Pt(10)

            if contact.get("linkedin") or contact.get("github"):
                links = []
                if contact.get("linkedin"):
                    links.append(contact["linkedin"])
                if contact.get("github"):
                    links.append(contact["github"])
                links_para = doc.add_paragraph(" | ".join(links))
                links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                links_para.runs[0].font.size = Pt(10)

            # Add summary
            if context.get("summary"):
                self._add_section_heading(doc, "Professional Summary", template_name)
                summary_para = doc.add_paragraph(context["summary"])
                summary_para.runs[0].font.size = Pt(11)

            # Add experiences
            if context.get("experiences"):
                self._add_section_heading(doc, "Work Experience", template_name)
                for exp in context["experiences"]:
                    # Job title and company
                    title_para = doc.add_paragraph()
                    title_run = title_para.add_run(exp["title"])
                    title_run.font.size = Pt(12)
                    title_run.font.bold = True

                    company_para = doc.add_paragraph(exp["company"])
                    company_para.runs[0].font.size = Pt(11)
                    company_para.runs[0].font.italic = True

                    # Dates
                    dates_para = doc.add_paragraph(exp["dates"])
                    dates_para.runs[0].font.size = Pt(10)

                    # Achievements
                    for ach in exp["achievements"]:
                        ach_para = doc.add_paragraph(ach["text"], style="List Bullet")
                        ach_para.runs[0].font.size = Pt(11)

                    # Technologies
                    if exp.get("technologies"):
                        tech_para = doc.add_paragraph()
                        tech_para.add_run("Technologies: ").font.bold = True
                        tech_para.add_run(", ".join(exp["technologies"]))
                        tech_para.runs[0].font.size = Pt(10)
                        tech_para.runs[1].font.size = Pt(10)

            # Add skills
            if context.get("skills_by_category"):
                self._add_section_heading(doc, "Skills", template_name)
                for category, skills in context["skills_by_category"].items():
                    cat_para = doc.add_paragraph()
                    cat_run = cat_para.add_run(f"{category}: ")
                    cat_run.font.bold = True
                    cat_run.font.size = Pt(11)

                    skill_names = [s["name"] for s in skills]
                    cat_para.add_run(", ".join(skill_names))
                    cat_para.runs[1].font.size = Pt(11)

            # Add education
            if context.get("education"):
                self._add_section_heading(doc, "Education", template_name)
                for edu in context["education"]:
                    degree_para = doc.add_paragraph()
                    degree_run = degree_para.add_run(edu["degree"])
                    degree_run.font.size = Pt(11)
                    degree_run.font.bold = True

                    inst_para = doc.add_paragraph(edu["institution"])
                    inst_para.runs[0].font.size = Pt(11)

                    if edu.get("graduation_year") or edu.get("gpa"):
                        details = []
                        if edu.get("graduation_year"):
                            details.append(edu["graduation_year"])
                        if edu.get("gpa"):
                            details.append(f"GPA: {edu['gpa']}")
                        details_para = doc.add_paragraph(" | ".join(details))
                        details_para.runs[0].font.size = Pt(10)

            # Add certifications
            if context.get("certifications"):
                self._add_section_heading(doc, "Certifications", template_name)
                for cert in context["certifications"]:
                    cert_text = f"{cert['name']} - {cert['issuer']}"
                    if cert.get("date"):
                        cert_text += f" ({cert['date']})"
                    cert_para = doc.add_paragraph(cert_text)
                    cert_para.runs[0].font.size = Pt(11)

            # Create output directory if needed
            output_path.parent.mkdir(parents=True, exist_ok=True)

            # Save document
            doc.save(str(output_path))

            logger.info(f"Generated DOCX: {output_path} using template '{template_name}'")

            return output_path

        except Exception as e:
            raise Exception(f"Failed to generate DOCX: {e}") from e

    def _add_section_heading(
        self, doc: Any, text: str, template_name: str
    ) -> None:
        """Add a section heading with template-specific styling."""
        from docx.shared import Pt, RGBColor

        heading = doc.add_paragraph(text)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True

        if template_name == "modern":
            heading_run.font.color.rgb = RGBColor(30, 64, 175)
        elif template_name == "classic":
            heading_run.font.all_caps = True

        heading.space_before = Pt(12)
        heading.space_after = Pt(6)
